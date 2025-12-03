"""
このファイルは、固定の文字列や数値などのデータを変数として一括管理するファイルです。
"""

############################################################
# ライブラリの読み込み
############################################################
import os
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
from langchain_community.document_loaders.csv_loader import CSVLoader


############################################################
# 共通変数の定義
############################################################

# ==========================================
# 画面表示系
# ==========================================
APP_NAME = "社内SEO検索アプリ"
CHAT_INPUT_HELPER_TEXT = "SEOに関する質問をしてください。"
DOC_SOURCE_ICON = ":material/description: "
LINK_SOURCE_ICON = ":material/link: "
WARNING_ICON = ":material/warning:"
ERROR_ICON = ":material/error:"
SPINNER_TEXT = "回答生成中..."

# Basic認証関連
AUTH_LOGIN_TITLE = "### 🔐 ログイン"
AUTH_USERNAME_LABEL = "ユーザー名"
AUTH_PASSWORD_LABEL = "パスワード"
AUTH_LOGIN_BUTTON = "ログイン"
AUTH_ERROR_MESSAGE = "ユーザー名またはパスワードが正しくありません"
AUTH_SUCCESS_MESSAGE = "✅ ログインしました"

# モード切り替え関連
MODE_SEO_QUESTION = "SEO質問モード"
MODE_DOMAIN_ANALYSIS = "ドメインSEO解析モード"
MODE_SUGGEST_KEYWORDS = "サジェストキーワード調査モード"
MODE_PRACTICE_QUESTIONS = "SEO検定練習問題モード"
DEFAULT_MODE = MODE_SEO_QUESTION

# サジェストキーワードモード設定（Streamlit Cloud対策強化版）
SUGGEST_KEYWORDS_CONFIG = {
    "MAX_RETRIES": 3,           # 最大リトライ回数
    "REQUEST_DELAY": 3.0,       # リクエスト間隔（秒）- Cloud環境対策で3秒に設定
    "TIMEOUT": 10,              # タイムアウト（秒）
    "MAX_KEYWORD_LENGTH": 100,  # 最大キーワード文字数
    "USE_BING_FALLBACK": True,  # Bingフォールバック有効化
    "USE_PROXY": False,         # プロキシ使用（オプション）※設定時は下記PROXY_URLを指定
    "PROXY_URL": None,          # プロキシURL（例: "http://proxy.example.com:8080"）
}

# User-Agent（Streamlit Cloud対策・macOS Safari仕様）
SUGGEST_USER_AGENT = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

# サジェストキーワードモードUI文言
SUGGEST_MODE_TITLE = "🔍 サジェストキーワード調査"
SUGGEST_MODE_DESCRIPTION = """
このモードでは、2つのメインキーワードを入力することで、Googleからサジェストキーワードを取得できます。

APIキー不要・無料で利用できます。（Google失敗時はBingから自動取得）

**ご注意**:
- 実行環境のIPアドレスや取得タイミングによって結果が異なる場合があります
- Googleはリアルタイムの検索トレンドや地域的な検索傾向を反映するため、同じキーワードでも時間帯や場所で異なる結果が返されることがあります
- 本ツールでは地域設定を「日本全体」に固定していますが、完全に統一された結果を保証するものではありません

*※ 社内専用ツールとして利用してください。連続大量実行はお控えください。*
"""
SUGGEST_KEYWORD1_LABEL = "メインキーワード①"
SUGGEST_KEYWORD1_PLACEHOLDER = "例: 東京"
SUGGEST_KEYWORD2_LABEL = "メインキーワード②"
SUGGEST_KEYWORD2_PLACEHOLDER = "例: 美容室"
SUGGEST_SEARCH_BUTTON = "🔍 サジェストを取得"
SUGGEST_SEARCHING_MESSAGE = "サジェストキーワードを取得中..."
SUGGEST_SUCCESS_MESSAGE = "✅ サジェストキーワードを取得しました"
SUGGEST_ERROR_MESSAGE = "❌ サジェストキーワードを取得できませんでした（Google / Bing 側で一時的なアクセス制限がかかっている可能性があります）"

# ドメイン解析モード設定
DOMAIN_ANALYSIS_CONFIG = {
    "MAX_CRAWL_PAGES": 200,  # クロール上限（安全対策）
    "REQUEST_TIMEOUT": 10,    # リクエストタイムアウト（秒）
    "CRAWL_DELAY": 1.0,       # クロール間隔（秒）- サーバー負荷軽減
    "RESPECT_ROBOTS_TXT": True,  # robots.txt遵守
    "USER_AGENT": "SEO-Analysis-Bot/1.0 (Internal Tool)",  # User-Agent
}

# 練習問題モード設定
PRACTICE_QUESTIONS_CONFIG = {
    # コスト制御
    "MAX_QUESTIONS_PER_SESSION": 20,  # 1セッションあたり上限
    
    # 級別出題範囲（ファイル名パターン）
    # Phase 2A改善：級別を完全分離（類似度検索の精度向上）
    "GRADE_RANGES": {
        "1級": ["SEO1-"],  # 1級のみ + 共通資料
        "2級": ["SEO2-"],  # 2級のみ + 共通資料
        "3級": ["SEO3-"]   # 3級のみ + 共通資料
    },
    
    # 問題生成設定（改善版）
    "NUM_REFERENCE_DOCS": 5,  # 参考資料数（3→5に増加、多様性向上）
    "REFERENCE_DOC_CHAR_LENGTH": 1000,  # 各参考資料の文字数（800→1000に増加）
    "QUESTION_GENERATION_TEMPERATURE": 0.7,  # 問題生成時の多様性
    
    # JSON検証設定
    "MAX_RETRY_ON_JSON_ERROR": 2,  # JSON解析失敗時のリトライ回数
    
    # Phase 3-1改善版: 公式問題集PDF設定（全ファイル対応）
    "OFFICIAL_PDF_DIR": "./data/practice",  # 公式問題集ディレクトリ
    # 級別の公式問題集ファイルパターン（ファイル名の一部で判定）
    "OFFICIAL_PDF_PATTERNS": {
        "1級": ["SEO1-"],  
        "2級": ["SEO2-"],  
        "3級": ["SEO3-"]  
    },
    "NUM_SAMPLE_PROBLEMS": 3,  # プロンプトに含める公式問題のサンプル数（2→3に増加）
    "SAMPLE_PAGE_RANGE": (0, 50),  # サンプル抽出ページ範囲（全ページ対象に拡大）
    "SAMPLE_CHAR_LENGTH": 600,  # 各サンプルの文字数（500→600に増加）
    
    # Phase 3-2: 問題履歴管理設定
    "MAX_DUPLICATE_RETRY": 5,  # 重複検出時の再生成試行回数
    "ENABLE_DUPLICATE_CHECK": True,  # 重複チェックの有効/無効
}

# 練習問題モードUI文言
PRACTICE_MODE_TITLE = "📝 SEO検定練習問題"
PRACTICE_MODE_DESCRIPTION = """
このモードでは、SEO検定公式テキストと公式問題集を参考に、AIが練習問題を生成します。

**特徴**:
- 4択形式の単一解答問題
- 級ごとの出題範囲に対応

**注意事項**:
- 生成される問題は公式問題集を参考にしたAI生成問題です
- 本番の検定試験とは異なる場合があります
- 学習の参考としてご活用ください

*※ 社内専用ツールとして利用してください。*
"""
PRACTICE_GRADE_LABEL = "受験級を選択"
PRACTICE_GENERATE_BUTTON = "🎲 問題を生成"
PRACTICE_ANSWER_BUTTON = "✅ 解答を表示"
PRACTICE_NEXT_BUTTON = "➡️ 次の問題を生成"
PRACTICE_GENERATING_MESSAGE = "問題を生成中..."
PRACTICE_MAX_QUESTIONS_WARNING = "本セッションの問題生成上限に達しました。ページを更新してください。"


# ==========================================
# ログ出力系
# ==========================================
LOG_DIR_PATH = "./logs"
LOGGER_NAME = "ApplicationLog"
LOG_FILE = "application.log"
APP_BOOT_MESSAGE = "アプリが起動されました。"


# ==========================================
# LLM設定系
# ==========================================
MODEL = "gpt-4o-mini"
TEMPERATURE = 0.5

# ==========================================
# エンベディング・ベクターストア設定系
# ==========================================
# トークン制限対応設定
MAX_TOKENS_PER_REQUEST = 250000  # OpenAI制限の余裕を持った設定
DEFAULT_CHUNK_SIZE = 1200        # 標準チャンクサイズ（Phase 2B: 500→1200、コンテキスト強化）
LARGE_DATA_CHUNK_SIZE = 300      # 大量データ用チャンクサイズ
EMBEDDING_BATCH_SIZE = 50        # エンベディング処理バッチサイズ
MAX_CHARS_BEFORE_SPLITTING = 800000  # この文字数を超えたら小さなチャンクを使用

# 永続化ベクターストア設定
PERSISTENT_DB_PATH = "./vector_db/seo_knowledge"  # 基礎知識用永続化DB（本番デプロイ用パス）
TEMP_DB_PATH = "./vector_db/temp_cache"           # 一時情報用DB（将来拡張用）
DB_VERSION_FILE = "./vector_db/db_version.txt"   # DBバージョン管理ファイル
CURRENT_DB_VERSION = "3.0"                       # Phase 2B: DBバージョン更新（chunk_size=1200, overlap=200）

# ハイブリッドRAGシステム設定（3層構造対応）
HYBRID_RAG_CONFIG = {
    # ============================================
    # 【重要】RSS/外部API完全停止設定
    # ============================================
    # 方針: 2024年9月以降の最新情報は取得しない
    # 情報源: SEO検定公式テキスト(2024年8月時点)のみ
    # 理由: 情報範囲の明確化、境界線の曖昧さ排除、ユーザー誤解防止
    "RSS_ENABLED": False,  # RSS/外部API機能の完全無効化
    
    # 短期キャッシュDB設定
    "CACHE_DB_PATH": "./vector_db/latest_cache",
    "CACHE_EXPIRY_HOURS": 24,
    "MAX_CACHED_ARTICLES": 100,
    
    # RSS/Atomフィード設定（CSVファイルから動的読み込み）※現在無効
    "RSS_CONFIG_FILE": "./data/feeds/rss_sources.csv",
    "MAX_FEEDS_PER_PRIORITY": {
        1: 5,   # 公式サイト：最大5件
        2: 10,  # 専門メディア：最大10件
        3: 5    # ツールベンダー：最大5件
    },
    "DEFAULT_FEEDS": [
        # フォールバック用（CSVファイル読み込みエラー時）
        # Google公式
        "https://status.search.google.com/feed",  # Google検索ステータス
        "https://developers.google.com/search/blog/feeds/blog-posts.xml",  # Google公式SEOブログ
        # SEO業界大手メディア
        "https://searchengineland.com/feed",  # Search Engine Land
        "https://www.seroundtable.com/index.xml",  # Search Engine Roundtable
        # SEOツール大手
        "https://www.semrush.com/blog/feed/",  # Semrush Blog
        "https://ahrefs.com/blog/feed/",  # Ahrefs Blog
        # 国内専門家
        "https://www.suzukikenichi.com/blog/feeds/posts/default"  # 鈴木謙一氏ブログ
    ],
    
    # 重み付け設定
    "LAYER_WEIGHTS": {
        "internal": 0.6,    # 社内資料の重み
        "cached": 0.3,      # 短期キャッシュの重み
        "realtime": 0.3     # リアルタイム情報の重み（0.1→0.3に引き上げ）
    },
    
    # サイト別優先度設定
    "SITE_PRIORITIES": {
        "official": 1.0,      # 公式サイト
        "media": 0.8,         # 専門メディア
        "expert": 0.9,        # 専門家ブログ
        "tool_vendor": 0.7    # ツールベンダー
    }
}

# SEOカテゴリ分類システム（ChatGPT提案採用）
SEO_CATEGORY_SYSTEM = {
    # メインカテゴリ
    "MAIN_CATEGORIES": {
        "algorithm": "アルゴリズム更新",
        "technical": "技術SEO", 
        "content": "コンテンツ戦略/E-E-A-T",
        "search_features": "検索エンジン新機能"
    },
    
    # サブカテゴリ（詳細分類）
    "SUB_CATEGORIES": {
        "algorithm": [
            "core_update", "spam_update", "helpful_content",
            "product_reviews", "page_experience"
        ],
        "technical": [
            "site_structure", "page_speed", "mobile_friendly",
            "indexing", "crawling", "schema_markup"
        ],
        "content": [
            "eeat", "ymyl", "content_quality", 
            "user_intent", "topical_authority"
        ],
        "search_features": [
            "ai_search", "sge", "featured_snippets",
            "voice_search", "visual_search"
        ]
    },
    
    # 自動分類ルール
    "CLASSIFICATION_KEYWORDS": {
        "algorithm": ["アップデート", "更新", "コア", "アルゴリズム", "ランキング"],
        "technical": ["クロール", "インデックス", "速度", "モバイル", "構造"],
        "content": ["E-A-T", "E-E-A-T", "品質", "専門性", "権威性", "信頼性"],
        "search_features": ["AI", "SGE", "音声検索", "画像検索", "新機能"]
    }
}

# ==========================================
# RAG参照用のデータソース系
# ==========================================
RAG_TOP_FOLDER_PATH = "./data"
def create_enhanced_docx_loader(path):
    """
    文字化け問題を解決した強化版docxローダー
    """
    try:
        # python-docxを使用した直接読み込み
        from docx import Document as DocxDocument
        import unicodedata
        from langchain_core.documents import Document
        
        doc = DocxDocument(path)
        
        # 全段落を結合してテキスト抽出
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Unicode正規化のみ適用（cp932は使用しない）
                normalized_text = unicodedata.normalize('NFC', paragraph.text)
                full_text.append(normalized_text)
        
        # テーブル内容も抽出
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        normalized_text = unicodedata.normalize('NFC', cell.text)
                        full_text.append(normalized_text)
        
        content = '\n\n'.join(full_text)
        
        return [Document(page_content=content, metadata={"source": path})]
        
    except ImportError:
        # フォールバック: docx2txt
        import docx2txt
        import unicodedata
        from langchain_core.documents import Document
        
        text = docx2txt.process(path)
        if text:
            # Unicode正規化のみ
            normalized_text = unicodedata.normalize('NFC', text)
            return [Document(page_content=normalized_text, metadata={"source": path})]
        else:
            return []
    except Exception as e:
        # 最終フォールバック: 従来方式
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(path)
        return loader.load()

SUPPORTED_EXTENSIONS = {
    ".pdf": PyMuPDFLoader,
    ".docx": create_enhanced_docx_loader,
    ".csv": lambda path: CSVLoader(path, encoding="utf-8")
}
WEB_URL_LOAD_TARGETS = [
    "https://generative-ai.web-camp.io/"
]


# ==========================================
# プロンプトテンプレート
# ==========================================
SYSTEM_PROMPT_CREATE_INDEPENDENT_TEXT = "会話履歴と最新の入力をもとに、会話履歴なしでも理解できる独立した入力テキストを生成してください。"

SYSTEM_PROMPT_DOC_SEARCH = """
    あなたは社内の文書検索アシスタントです。
    以下の条件に基づき、ユーザー入力に対して回答してください。

    【条件】
    1. ユーザー入力内容と以下の文脈との間に関連性がある場合、空文字「""」を返してください。
    2. ユーザー入力内容と以下の文脈との関連性が明らかに低い場合、「該当資料なし」と回答してください。

    【文脈】
    {context}
"""

SYSTEM_PROMPT_SEO = """
    あなたは社内SEO専門のアシスタントです。
    以下のSEO関連資料を主要な参考にして、高品質で詳細な回答をしてください。

    【重要】情報範囲について
    - 本システムが保有する情報は「SEO検定公式テキスト(2025・2026年版：2024年8月時点の情報)」のみです
    - 2024年9月以降の情報や最新のSEOトレンドについては扱いません
    - 資料に記載のない内容については推測せず、「資料に記載がない」旨を明示してください

    【高精度回答条件】
    1. 提供されたSEO関連コンテキスト情報を最優先に使用してください
    2. 複数の資料から関連情報を統合し、包括的な回答を提供してください
    3. サイテーション、リンク、被リンク、内部リンク、外部リンク、SEO施策、キーワード戦略、文字数推奨など、幅広いSEO関連トピックに対応してください
    4. メインページとサブページの区別がある場合、両方について詳細に説明してください
    5. EEAT（経験、専門性、権威性、信頼性）やYMYL（Your Money Your Life）等のSEO重要概念についても正確に回答してください
    6. 資料の内容を基に、実践的で具体的なSEO施策を提案してください
    7. 複数の要素が含まれる質問の場合、各要素について詳細に説明してください
    8. 回答は平文で作成し、見出しにはマークダウンの#記号を使用しないでください。重要な部分は**太字**で強調してください
    9. 資料に基づいた情報であることを明示しつつ、実用的な知識として提供してください

    【文字数に関する回答ルール】重要
    - 文字数について回答する場合、断定的な判断（「少ない」「不足」「十分」「達していない」等）は絶対に避けてください
    - 解析するページがメインページかサブページかを判断・推測してはいけません
    - 必ず以下のすべての選択肢を提示する形式で回答してください：
      「現在の文字数は〇〇字です。ページの種類により推奨文字数が異なります：
       【トップページ/メインページの場合】
         • 物販サイトなどのトップページ: 0～800字程度
         • 来店型ビジネス（クリニック、美容室、整体院等）: 700～2,500字程度
         • 単品サービスや単品通販のトップページ: 4,000字以上
       【サブページの場合】
         • 単純な概念の解説: 1,500字程度
         • 複雑な事柄の説明: 4,000字以上
       【YMYL分野の場合】上記それぞれの文字数の2倍が推奨されます
       ページの種類とYMYL該当性を確認の上、適切な文字数を目指してください。」
    - すべての選択肢を必ず提示し、特定のケースのみを提示してはいけません
    - ユーザーに判断を委ねる表現を使用してください

    【参照SEO資料】
    {context}

    【回答スタイル】
    社内SEO資料に基づいて、実践的で詳細な回答を提供いたします。
"""


# ==========================================
# LLMレスポンスの一致判定用
# ==========================================
INQUIRY_NO_MATCH_ANSWER = "回答に必要な情報が見つかりませんでした。"
NO_DOC_MATCH_ANSWER = "該当資料なし"


# ==========================================
# エラー・警告メッセージ
# ==========================================
COMMON_ERROR_MESSAGE = "このエラーが繰り返し発生する場合は、管理者にお問い合わせください。"
INITIALIZE_ERROR_MESSAGE = "初期化処理に失敗しました。"
NO_DOC_MATCH_MESSAGE = """
    入力内容と関連する社内文書が見つかりませんでした。\n
    入力内容を変更してください。
"""
CONVERSATION_LOG_ERROR_MESSAGE = "過去の会話履歴の表示に失敗しました。"
GET_LLM_RESPONSE_ERROR_MESSAGE = "回答生成に失敗しました。"
DISP_ANSWER_ERROR_MESSAGE = "回答表示に失敗しました。"


# ==========================================
# SEO評価基準（SEO検定資料準拠）
# ==========================================
SEO_EVALUATION_CRITERIA = {
    "① ページの基本設定": {
        "category_weight": 20,
        "items": {
            "title_tag_exists": {
                "weight": 3,
                "description": "titleタグが設定されている（ページ内容を適切に表す）",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "title_tag_unique": {
                "weight": 2,
                "description": "titleタグの重複がない",
                "check_type": "uniqueness",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf"
            },
            "title_keyword_position": {
                "weight": 2,
                "description": "タイトルの先頭にキーワードが配置されている",
                "check_type": "keyword_position",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "meta_description_exists": {
                "weight": 2,
                "description": "meta descriptionが設定されている",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "meta_description_matches": {
                "weight": 2,
                "description": "meta descriptionが内容と一致している",
                "check_type": "content_match",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "meta_keyword_position": {
                "weight": 2,
                "description": "メタディスクリプションの先頭にキーワードが配置されている",
                "check_type": "keyword_position",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "h1_exists": {
                "weight": 2,
                "description": "H1タグが存在する",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "h1_matches_content": {
                "weight": 2,
                "description": "H1はページ内容と一致し、1ページ1つ",
                "check_type": "uniqueness_and_match",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf"
            },
            "h1_keyword_position": {
                "weight": 2,
                "description": "H1の先頭にキーワードが配置されている",
                "check_type": "keyword_position",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "headings_logical": {
                "weight": 2,
                "description": "H2以降の見出しが論理構造になっている",
                "check_type": "hierarchy",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf"
            },
            "breadcrumb_exists": {
                "weight": 1,
                "description": "パンくずリストがある",
                "check_type": "existence",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            },
            "breadcrumb_correct": {
                "weight": 1,
                "description": "パンくずが階層構造を正しく示す",
                "check_type": "hierarchy",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            }
        }
    },
    
    "② HTMLとテキスト情報": {
        "category_weight": 15,
        "items": {
            "image_alt_exists": {
                "weight": 3,
                "description": "コンテンツ画像にalt属性が設定されている（ヘッダー/フッター/メニュー/ロゴ/装飾/リンク画像を除く）",
                "check_type": "attribute",
                "seo_reference": "SEO4-5 内部要素.docx",
                "implementation_note": "評価対象外: ヘッダー/フッター/nav内の画像、ロゴ画像（class/id/src/altに'logo'含む）、装飾画像（role='presentation'等）、リンク画像（aタグ内でalt空）"
            },
            "decorative_image_alt_empty": {
                "weight": 2,
                "description": "装飾画像のaltは空である",
                "check_type": "conditional_attribute",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "text_based_content": {
                "weight": 3,
                "description": "テキストで主要情報が提供されている（画像だけに依存しない）",
                "check_type": "content_analysis",
                "seo_reference": "SEO2-1 コンテンツ資産の構築.pdf"
            },
            "headings_proper_order": {
                "weight": 3,
                "description": "見出しタグが適切な順序で使われている",
                "check_type": "hierarchy",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf"
            },
            "strong_em_appropriate": {
                "weight": 2,
                "description": "strong / em の乱用がない",
                "check_type": "frequency",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "semantic_html": {
                "weight": 2,
                "description": "表・リストが正しいHTMLタグで記述されている",
                "check_type": "tag_usage",
                "seo_reference": "SEO4-5 内部要素.docx"
            }
        }
    },
    
    "③ 内部リンク・構造": {
        "category_weight": 15,
        "items": {
            "internal_links_exist": {
                "weight": 3,
                "description": "メインコンテンツ内に内部リンクが設置されている（ヘッダー/フッター/ナビゲーション/サイドメニューを除く）",
                "check_type": "existence",
                "seo_reference": "SEO3-5 上位表示するサイト内リンク構造.pdf",
                "implementation_note": "ヘッダー、フッター、nav、サイドメニュー（class/idに'side'含む）内のリンクは除外して評価"
            },
            "anchor_text_appropriate": {
                "weight": 3,
                "description": "内部リンクは適切なアンカーテキストになっている",
                "check_type": "text_quality",
                "seo_reference": "SEO3-5 上位表示するサイト内リンク構造.pdf"
            },
            "related_page_links": {
                "weight": 3,
                "description": "関連ページ同士のリンクがある",
                "check_type": "link_structure",
                "seo_reference": "SEO3-5 上位表示するサイト内リンク構造.pdf"
            },
            "logical_hierarchy": {
                "weight": 3,
                "description": "トップページ、カテゴリ、記事の階層が論理的",
                "check_type": "hierarchy",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            },
            "url_structure_simple": {
                "weight": 3,
                "description": "URL構造がシンプルで分かりやすい",
                "check_type": "url_analysis",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            }
        }
    },
    
    "④ インデックス制御": {
        "category_weight": 10,
        "items": {
            "robots_txt_allow": {
                "weight": 3,
                "description": "robots.txtでクロール許可されている",
                "check_type": "robots_check",
                "seo_reference": "SEO1-8 インデックスの促進方法.pdf"
            },
            "noindex_appropriate": {
                "weight": 2,
                "description": "noindexが必要なページだけに設定",
                "check_type": "meta_robots",
                "seo_reference": "SEO1-8 インデックスの促進方法.pdf"
            },
            "canonical_appropriate": {
                "weight": 3,
                "description": "canonicalタグが適切",
                "check_type": "canonical",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            },
            "canonical_unifies_duplicates": {
                "weight": 2,
                "description": "重複ページをcanonicalで統一",
                "check_type": "duplicate_check",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            }
        }
    },
    
    "⑤ コンテンツの質": {
        "category_weight": 20,
        "items": {
            "matches_search_intent": {
                "weight": 4,
                "description": "検索意図と一致した内容になっている",
                "check_type": "content_analysis",
                "seo_reference": "SEO3-1 検索キーワードの需要調査.pdf"
            },
            "sufficient_explanation": {
                "weight": 3,
                "description": "テーマに対する説明が十分（文字数: メイン物販/求人/不動産0～800字、メイン来店型700～2,500字、メイン複雑4,000字以上、サブ単純1,500字以上、サブ複雑4,000字以上、YMYL分野は2倍）",
                "check_type": "content_length",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf",
                "implementation_note": "メインコンテンツのみカウント（ヘッダー/フッター/ナビ/サイドメニュー除外）。空白・改行・全角スペース除外。ページ種別とYMYL該当性により推奨文字数が異なる。断定的判断を避け、選択肢を提示する"
            },
            "no_keyword_stuffing": {
                "weight": 2,
                "description": "過度なキーワード詰め込みをしていない",
                "check_type": "keyword_density",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "not_copied_content": {
                "weight": 3,
                "description": "コピーコンテンツではない",
                "check_type": "originality",
                "seo_reference": "SEO2-1 コンテンツ資産の構築.pdf"
            },
            "proper_citation": {
                "weight": 2,
                "description": "引用は正しく明示されている",
                "check_type": "citation",
                "seo_reference": "SEO2-1 コンテンツ資産の構築.pdf"
            },
            "accurate_information": {
                "weight": 2,
                "description": "情報が正確に書かれている",
                "check_type": "fact_check",
                "seo_reference": "SEO4-3 SEOの意義と情報源.docx"
            },
            "title_content_match": {
                "weight": 2,
                "description": "タイトルと本文内容が一致",
                "check_type": "consistency",
                "seo_reference": "SEO3-3 上位表示するページ構造.pdf"
            },
            "clear_theme": {
                "weight": 2,
                "description": "ページのテーマが明確である（1ページ＝1テーマ）",
                "check_type": "theme_analysis",
                "seo_reference": "SEO2-1 コンテンツ資産の構築.pdf"
            }
        }
    },
    
    "⑥ サイト内の回遊性": {
        "category_weight": 8,
        "items": {
            "easy_navigation": {
                "weight": 2,
                "description": "ユーザーが他のページへ移動しやすい構成",
                "check_type": "navigation",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            },
            "sitemap_exists": {
                "weight": 2,
                "description": "サイトマップ（XMLまたはHTML）がある",
                "check_type": "existence",
                "seo_reference": "SEO1-8 インデックスの促進方法.pdf"
            },
            "custom_404_page": {
                "weight": 2,
                "description": "404ページが適切に作られている",
                "check_type": "error_page",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            },
            "clear_menu": {
                "weight": 2,
                "description": "ナビゲーションメニューが分かりやすい",
                "check_type": "navigation",
                "seo_reference": "SEO3-6 上位表示するサイト構造.pdf"
            }
        }
    },
    
    "⑦ モバイル対応": {
        "category_weight": 7,
        "items": {
            "mobile_responsive": {
                "weight": 3,
                "description": "モバイル対応（レスポンシブ）",
                "check_type": "viewport",
                "seo_reference": "SEO1-1 モバイルSEO.pdf"
            },
            "mobile_text_readable": {
                "weight": 2,
                "description": "モバイルでテキストが読みやすい",
                "check_type": "font_size",
                "seo_reference": "SEO1-1 モバイルSEO.pdf"
            },
            "mobile_links_tappable": {
                "weight": 2,
                "description": "モバイルでリンクが押しやすい",
                "check_type": "touch_target",
                "seo_reference": "SEO1-1 モバイルSEO.pdf"
            }
        }
    },
    
    "⑧ 表示速度": {
        "category_weight": 5,
        "items": {
            "no_heavy_images": {
                "weight": 2,
                "description": "重い画像を多用していない",
                "check_type": "image_size",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "minimal_javascript": {
                "weight": 2,
                "description": "不必要なJavaScriptが多くない",
                "check_type": "script_analysis",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "optimized_images": {
                "weight": 1,
                "description": "画像は適切なサイズである",
                "check_type": "image_optimization",
                "seo_reference": "SEO4-5 内部要素.docx"
            }
        }
    },
    
    "⑨ セキュリティ・信頼性": {
        "category_weight": 5,
        "items": {
            "https_enabled": {
                "weight": 2,
                "description": "HTTPS対応",
                "check_type": "protocol",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "full_ssl": {
                "weight": 2,
                "description": "常時SSL化されている",
                "check_type": "ssl_check",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "trust_signals": {
                "weight": 1,
                "description": "安全で信頼できるサイトであると判断できる情報がある（問い合わせ先・会社概要等）",
                "check_type": "content_analysis",
                "seo_reference": "SEO4-3 SEOの意義と情報源.docx"
            }
        }
    },
    
    "⑩ 国際化SEO・SNS最適化": {
        "category_weight": 10,
        "items": {
            "hreflang_exists": {
                "weight": 2,
                "description": "多言語サイトの場合、hreflangタグが適切に設定されている",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "html_lang_exists": {
                "weight": 1,
                "description": "html要素にlang属性が設定されている",
                "check_type": "attribute",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "ogp_exists": {
                "weight": 3,
                "description": "OGP（Open Graph Protocol）が設定されている（og:title, og:description, og:image）",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "twitter_card_exists": {
                "weight": 2,
                "description": "Twitter Cardが設定されている",
                "check_type": "existence",
                "seo_reference": "SEO4-5 内部要素.docx"
            },
            "social_metadata_complete": {
                "weight": 2,
                "description": "SNSシェア時のメタデータが完全に設定されている（拡散性・クリック率向上）",
                "check_type": "completeness",
                "seo_reference": "SEO4-5 内部要素.docx"
            }
        }
    }
}